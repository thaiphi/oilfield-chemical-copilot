from __future__ import annotations

from pathlib import Path

from docx import Document

from oilfield_chemical_copilot.ingest.chunking import chunk_text
from oilfield_chemical_copilot.ingest.models import LoadedChunk


def parse_docx(path: Path, *, source_root: Path | None = None, topic: str = "unknown") -> list[LoadedChunk]:
    document = Document(path)
    parts: list[str] = []
    parts.extend(paragraph.text for paragraph in document.paragraphs if paragraph.text.strip())
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))
    return chunk_text(
        text="\n".join(parts),
        source_file=path,
        source_root=source_root,
        topic=topic,
        parser_type="docx",
        page_or_sheet="document",
    )
